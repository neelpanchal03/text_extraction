from django.shortcuts import render, redirect
from .models import Document
from .tasks import extract_text_async
from asgiref.sync import sync_to_async
import requests
import tempfile
from django.core.files.base import ContentFile
from urllib.parse import urlparse
from pathlib import Path
import os
import fitz
import docx
from django.core.exceptions import ValidationError



async def upload_file(request):
    if request.method == 'POST':
        file = request.FILES.get('file')
        email = request.POST.get('email')

        if not file or not email:
            return render(request, 'upload_file.html', {
                'error': 'Both email and file are required.'
            })

        if file:
            valid_extensions = ['pdf', 'doc', 'docx']
            valid_content_types = ['application/pdf',
                                   'application/msword',
                                   'application/vnd.openxmlformats-officedocument.wordprocessingml.document']

            file_extension = file.name.split('.')[-1].lower()
            if file_extension not in valid_extensions or file.content_type not in valid_content_types:
                return render(request, 'upload_file.html', {
                    'error': 'Unsupported file type. Please upload a PDF, DOC, or DOCX file.'
                })

            doc = await sync_to_async(Document.objects.create)(file=file, email=email)
            extract_text_async.delay(doc.id)

            return render(request, 'upload_file.html', {
                'success': 'Your file has been uploaded successfully and is being processed!'
            })

    return render(request, 'upload_file.html')


def download_google_file(file_url, destination, export_format="pdf"):
    print(file_url,'file_url')
    try:
        if "drive.google.com" in file_url:
            file_id = file_url.split("/d/")[1].split("/")[0]
            url = f"https://drive.google.com/uc?export=download&id={file_id}"
        elif "docs.google.com" in file_url:
            file_id = file_url.split("/d/")[1].split("/")[0]
            url = f"https://docs.google.com/document/d/{file_id}/export?format={export_format}"
        else:
            url = file_url

        response = requests.get(url)

        print(response,'aagcggchsg')
        if response.status_code == 200:
            with open(destination, "wb") as file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        file.write(chunk)
            return destination, None
        else:
            return None, f"Failed to download the file. Status code: {response.status_code}"

    except Exception as e:
        return None, f"An error occurred: {e}"



async def upload_url(request):
    if request.method == 'POST':
        url = request.POST.get('url')
        email = request.POST.get('email')

        if url and is_valid_url(url):
            file_extension = url.split('.')[-1]
            suffix = f'.{file_extension}' if file_extension in ['pdf', 'docx', 'doc'] else '.pdf'

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                downloaded_file_path, error = download_google_file(url, temp_file.name)
                if error:
                    return render(request, 'upload_url.html', {'error': error})

                if not is_file_openable(downloaded_file_path):
                    return render(request, 'upload_url.html', {
                        'error': 'The downloaded file is corrupted or cannot be opened.'
                    })

                with open(downloaded_file_path, 'rb') as downloaded_file:
                    file_content = ContentFile(downloaded_file.read(), name=Path(downloaded_file_path).name)
                    doc = await sync_to_async(Document.objects.create)(file=file_content,file_public_url=url, email=email)
                    extract_text_async.delay(doc.id)

                os.unlink(downloaded_file_path)

            return render(request, 'upload_file.html', {
                'success': 'Your file has been uploaded successfully and is being processed!'
            })

    return render(request, 'upload_url.html')

def is_file_openable(file_path):
    try:
        if file_path.endswith('.pdf'):
            with fitz.open(file_path) as pdf:
                pdf[0].get_text("text")
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            if not doc.paragraphs:
                raise ValidationError("Document is empty.")
        elif file_path.endswith('.doc'):
            doc = docx.Document(file_path)
            if not doc.paragraphs:
                raise ValidationError("Document is empty.")

        return True
    except Exception as e:
        print(f"Error opening file {file_path}: {e}")
        return False

def is_valid_url(url):
    parsed_url = urlparse(url)
    return all([parsed_url.scheme, parsed_url.netloc])